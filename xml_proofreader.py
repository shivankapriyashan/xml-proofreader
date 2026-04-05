import os
import sys
import json
import argparse
import time
import logging
from pathlib import Path
from lxml import etree
from openai import OpenAI
from docx import Document
import psutil
from dotenv import load_dotenv
import langcodes
import numpy as np
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.docstore.document import Document as LangchainDocument

load_dotenv()

logs_dir = Path(__file__).parent / 'logs'
logs_dir.mkdir(exist_ok=True)
log_filename = logs_dir / f"xml-proofreader_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[logging.FileHandler(log_filename), logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

logging.getLogger("openai").setLevel(logging.WARNING)
logging.getLogger("httpx").setLevel(logging.WARNING)

def load_prompt_template():
    prompt_file = Path(__file__).parent / 'data' /'system_prompt_template.txt'
    with open(prompt_file, 'r', encoding='utf-8') as f:
        return f.read()

class XMLProofreader:
    def __init__(self, api_key, model, language="en", style_guide_path=None):
        self.client = OpenAI(api_key=api_key)
        self.model = model
        self.language = language
        self.api_key = api_key
        
        vector_store_path = Path(__file__).parent / 'vector_store'
        
        if style_guide_path:
            logger.info("Creating vector store from style guide...")
            self.vector_store = self._load_style_guide(style_guide_path)
            logger.debug(f"Created vector store with {len(self.vector_store.docstore._dict)} chunks")
        else:
            if vector_store_path.exists():
                logger.info("Loading cached vector store...")
                try:
                    embeddings = OpenAIEmbeddings(openai_api_key=api_key)
                    self.vector_store = FAISS.load_local(str(vector_store_path), embeddings)
                    logger.debug(f"Loaded vector store with {len(self.vector_store.docstore._dict)} chunks")
                except Exception as e:
                    raise ValueError(f"Failed to load vector store: {e}. Please provide --style-guide to create it.")
            else:
                raise ValueError("No vector store found. Please run with --style-guide first to create it.")
        
        self.prompt_template = load_prompt_template()
        self.start_time = None
        self.process = psutil.Process()
        self.initial_memory = None

    def _load_style_guide(self, docx_path):
        logger.debug(f"Loading style guide from: {docx_path}")
        doc = Document(docx_path)
        
        full_text = "\n\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
        
        if not full_text:
            raise ValueError("Style guide file is empty")
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=600,
            chunk_overlap=100,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        chunks = text_splitter.split_text(full_text)
        logger.debug(f"Split style guide into {len(chunks)} chunks")
        
        documents = [LangchainDocument(page_content=chunk) for chunk in chunks]
        
        embeddings = OpenAIEmbeddings(openai_api_key=self.api_key)
        vector_store = FAISS.from_documents(documents, embeddings)
        
        vector_store_path = Path(__file__).parent / 'vector_store'
        vector_store.save_local(str(vector_store_path))
        logger.debug(f"Saved vector store to {vector_store_path}")
        
        return vector_store

    def _get_relevant_rules(self, text, top_k=10):
            try:
                docs_and_scores = self.vector_store.similarity_search_with_score(text, k=top_k)
                
                logger.debug(f"Retrieved {len(docs_and_scores)} relevant rules")
                
                formatted_rules = "\n".join([
                    f"- {doc.page_content}"
                    for doc, score in docs_and_scores
                ])
                
                return formatted_rules
            except Exception as e:
                logger.warning(f"Error retrieving relevant rules: {e}")
                return ""

    def proofread_text(self, text):
        if not text.strip():
            return []
        
        try:
            relevant_rules = self._get_relevant_rules(text, top_k=5)
            
            try:
                lang_obj = langcodes.Language.get(self.language)
                lang_instruction = f'Proofread in {lang_obj.display_name("en")}.'
            except:
                lang_instruction = f'Proofread in {self.language} language.'
            
            dynamic_prompt = self.prompt_template.replace('{LANG_INSTRUCTION}', lang_instruction)
            dynamic_prompt = dynamic_prompt.replace('{STYLE_RULES}', relevant_rules)
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": dynamic_prompt},
                    {"role": "user", "content": f'Proofread this text (language: {self.language}):\n\n"{text}"'}
                ],
                temperature=0.1,
                response_format={"type": "json_object"}
            )
            
            data = json.loads(response.choices[0].message.content)
            errors = data.get("errors", [])
            errors = [e for e in errors if e.get('incorrect_text') != e.get('correct_text')]
            
            for error in errors:
                if not error.get("reason"):
                    error["reason"] = f"{error.get('error_type', 'Error')} detected"
            
            return errors
        except Exception as e:
            logger.error(f"Error calling API: {e}")
            return []

    def create_error_tag(self, error):
        error_type = error.get("error_type", "")
        correct = error.get("correct_text", "").replace('"', '&quot;')
        incorrect = error.get("incorrect_text", "")
        reason = error.get("reason", "Error detected").replace('"', '&quot;')
        return f'<error type="{error_type}" correction="{correct}" reason="{reason}">{incorrect}</error>'
    
    def annotate_text(self, text, errors):
        if not errors:
            return text
        
        annotated = text
        for error in reversed(errors):
            incorrect = error.get("incorrect_text", "")
            if incorrect in annotated:
                annotated = annotated.replace(incorrect, self.create_error_tag(error), 1)
        return annotated
    
    def process_xml_file(self, input_path, output_path):
        self.start_time = time.time()
        self.initial_memory = self.process.memory_info().rss / 1024 / 1024
        
        logger.info(f"{'='*60}")
        logger.info(f"Processing: {input_path}")
        logger.info(f"Language: {self.language}")
        logger.info(f"Model: {self.model}")
        logger.info(f"{'='*60}")
        
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                original_xml = f.read()
            parser = etree.XMLParser(remove_blank_text=False, strip_cdata=False)
            tree = etree.parse(input_path, parser)
            root = tree.getroot()
        except Exception as e:
            logger.error(f"Error reading/parsing XML: {e}")
            return False
        
        paragraphs = [elem for elem in root.iter() if etree.QName(elem).localname == 'p']
        logger.info(f"Found {len(paragraphs)} paragraph(s) to proofread")
        
        if not paragraphs:
            logger.warning("No <p> elements found!")
            return False
        
        total_errors = 0
        paragraphs_with_errors = 0
        
        def process_paragraph(i, p_elem):
            original_text = self._get_text(p_elem)
            logger.info(f"== Paragraph {i}/{len(paragraphs)}: {original_text[:80]}...")
            
            if not original_text.strip():
                return i, p_elem, [], original_text
            
            errors = self.proofread_text(original_text)
            logger.debug(f"  Found {len(errors)} error(s)")
            
            if errors:
                for err in errors:
                    logger.debug(f"    - {err['error_type']}: '{err['incorrect_text']}' -> '{err['correct_text']}'")
                    logger.debug(f"      Reason: {err.get('reason', 'N/A')}")
            
            return i, p_elem, errors, original_text
        
        with ThreadPoolExecutor(max_workers=5) as executor:
            futures = {executor.submit(process_paragraph, i, p_elem): (i, p_elem) 
                      for i, p_elem in enumerate(paragraphs, 1)}
            
            for future in as_completed(futures):
                try:
                    i, p_elem, errors, original_text = future.result()
                    
                    if errors:
                        paragraphs_with_errors += 1
                        annotated = self.annotate_text(original_text, errors)
                        
                        if not self._validate_text_length(original_text, annotated):
                            logger.error(f"Text-length invariant violated for paragraph {i}")
                            return False
                        
                        self._update_element(p_elem, annotated)
                        total_errors += len(errors)
                except Exception as e:
                    logger.error(f"Error processing paragraph: {e}")
                    return False
        
        logger.info(f"{'='*60}")
        logger.info(f"[SUCCESS] Total errors found: {total_errors}")
        logger.info(f"[SUCCESS] Paragraphs with errors: {paragraphs_with_errors}/{len(paragraphs)}")
        
        try:
            output_xml = etree.tostring(root, encoding='unicode', method='xml')
            output_xml = '<?xml version="1.0" encoding="UTF-8"?>\n' + output_xml
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(output_xml)
            
            logger.info(f"[SUCCESS] Output written to: {output_path}")
            
            elapsed = time.time() - self.start_time
            final_mem = self.process.memory_info().rss / 1024 / 1024
            
            logger.info(f"{'='*60}")
            logger.info("PERFORMANCE METRICS:")
            logger.info(f"  Total processing time: {elapsed:.2f} seconds")
            logger.info(f"  Memory used: {final_mem - self.initial_memory:.2f} MB")
            logger.info(f"  Peak memory: {final_mem:.2f} MB")
            logger.info(f"{'='*60}")
            
            return True
        except Exception as e:
            logger.error(f"Error writing output: {e}")
            return False
    
    def _get_text(self, element):
        text = element.text or ""
        for child in element:
            if child.text:
                text += child.text
            if child.tail:
                text += child.tail
        return text
    
    def _update_element(self, element, new_content):
        element.text = None
        for child in list(element):
            element.remove(child)
        
        try:
            temp = etree.fromstring(f"<temp>{new_content}</temp>")
            element.text = temp.text
            for child in temp:
                element.append(child)
        except:
            element.text = new_content
    
    def _strip_error_tags(self, text):
        try:
            temp = etree.fromstring(f"<temp>{text}</temp>")
            return ''.join(temp.itertext())
        except:
            import re
            return re.sub(r'<error[^>]*>|</error>', '', text)
    
    def _validate_text_length(self, original, annotated):
        if len(original) != len(self._strip_error_tags(annotated)):
            return False
        return True            

def main():
    parser = argparse.ArgumentParser(
        description="Proofread XML files using AI and inject error annotations",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""Examples:
  python xml_proofreader.py --input input.xml --lang en --style-guide StyleGuide.docx
        """
    )

    parser.add_argument("--input", required=True, help="Input XML file path")
    parser.add_argument("--lang", required=True, help="Language code")
    parser.add_argument("--style-guide", help="Path to style guide .docx file (required)")
    parser.add_argument("--warning", "-w", action="store_true", help="Enable WARNING level logging")
    parser.add_argument("--info", "-i", action="store_true", help="Enable INFO level logging (progress updates)")
    parser.add_argument("--verbose", "-v", action="store_true", help="Enable verbose logging (DEBUG level)")
    
    
    args = parser.parse_args()
    
    if args.verbose:
        logger.setLevel(logging.DEBUG)
    elif args.info:
        logger.setLevel(logging.INFO)
    elif args.warning:
        logger.setLevel(logging.WARNING)
    
    api_key = os.getenv("API_KEY")
    if not api_key:
        logger.error("Error: API_KEY environment variable not set")
        sys.exit(1)
    
    model = os.getenv("MODEL")
    logger.debug(f"Using model: {model}")
    
    if not os.path.exists(args.input):
        logger.error(f"Error: Input file not found: {args.input}")
        sys.exit(1)
    
    if args.style_guide and not os.path.exists(args.style_guide):
        logger.error(f"Error: Style guide file not found: {args.style_guide}")
        sys.exit(1)
    
    input_path = Path(args.input)
    output_path = input_path.parent / f"{input_path.stem}.corrected.xml"
    logger.debug(f"Output will be saved to: {output_path}")
    
    proofreader = XMLProofreader(
        api_key=api_key,
        model=model,
        language=args.lang,
        style_guide_path=args.style_guide
    )
    
    success = proofreader.process_xml_file(args.input, str(output_path))
    logger.info("\n[SUCCESS] Processing completed successfully!" if success else "\n[FAILED] Processing failed!")
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()